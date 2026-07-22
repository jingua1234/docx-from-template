#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
模板结构检查脚本（docx-from-template Skill）

作用：
  解析参考 Word（.docx）或 Excel（.xlsx）模板，输出标题层级、表格列头、
  工作表结构等摘要，供 Agent 在仿写交付文档前对齐版式。

执行原理：
  - .docx：用 python-docx 遍历段落与表格；Heading 样式视为章节标题。
  - .xlsx：用 openpyxl 只读模式读取各工作表首行作为列头，并抽样前几行。
  - 结果打印到 stdout（UTF-8），也可重定向到临时文本文件。

用法：
  python scripts/inspect_docx.py path/to/template.docx
  python scripts/inspect_docx.py path/to/template.xlsx
"""

from __future__ import annotations

import json
import sys
from pathlib import Path


def _ensure_utf8_stdout() -> None:
    """Windows 控制台下尽量将 stdout 设为 UTF-8，避免中文乱码。"""
    try:
        sys.stdout.reconfigure(encoding="utf-8")
    except Exception:
        pass


def inspect_docx(path: Path) -> dict:
    """
    解析 Word 文档结构。

    返回结构化 dict，同时由调用方负责打印人类可读报告。
    核心字段：
      - headings: [{level, style, text}]
      - paragraphs_sample: 非空段落前若干条预览
      - tables: [{index, rows, cols, headers, sample_row}]
    """
    try:
        from docx import Document
    except ImportError:
        print("错误: 未安装 python-docx，请执行: python -m pip install python-docx")
        sys.exit(2)

    doc = Document(str(path))
    headings = []
    para_sample = []

    for para in doc.paragraphs:
        text = (para.text or "").strip()
        style = para.style.name if para.style else ""
        # Heading 1 / 标题 1 等均视为标题
        is_heading = style.startswith("Heading") or style.startswith("标题")
        if is_heading and text:
            # 从样式名末尾数字推断级别，失败则记 0
            level = 0
            for ch in reversed(style):
                if ch.isdigit():
                    level = int(ch)
                    break
            headings.append({"level": level, "style": style, "text": text})
        elif text and len(para_sample) < 40:
            para_sample.append({"style": style, "text": text[:120]})

    tables = []
    for i, table in enumerate(doc.tables):
        rows = len(table.rows)
        cols = len(table.columns) if table.rows else 0
        headers = []
        sample_row = []
        if rows > 0:
            headers = [(c.text or "").strip() for c in table.rows[0].cells]
        if rows > 1:
            sample_row = [(c.text or "").strip()[:80] for c in table.rows[1].cells]
        tables.append(
            {
                "index": i,
                "rows": rows,
                "cols": cols,
                "headers": headers,
                "sample_row": sample_row,
            }
        )

    return {
        "type": "docx",
        "path": str(path),
        "heading_count": len(headings),
        "table_count": len(tables),
        "headings": headings,
        "paragraphs_sample": para_sample,
        "tables": tables,
    }


def inspect_xlsx(path: Path) -> dict:
    """
    解析 Excel 工作簿结构。

    每个 sheet 记录列头（首行）与最多 3 行样例，便于对齐 Bug 跟踪表等模板。
    """
    try:
        from openpyxl import load_workbook
    except ImportError:
        print("错误: 未安装 openpyxl，请执行: python -m pip install openpyxl")
        sys.exit(2)

    # data_only=False：保留公式文本，便于发现模板占位
    wb = load_workbook(str(path), read_only=True, data_only=False)
    sheets = []
    for name in wb.sheetnames:
        ws = wb[name]
        rows_iter = ws.iter_rows(values_only=True)
        headers = []
        samples = []
        try:
            first = next(rows_iter)
            headers = [("" if v is None else str(v)).strip() for v in first]
        except StopIteration:
            pass
        for _, row in zip(range(3), rows_iter):
            samples.append([("" if v is None else str(v)).strip()[:80] for v in row])
        sheets.append({"name": name, "headers": headers, "sample_rows": samples})
    wb.close()

    return {
        "type": "xlsx",
        "path": str(path),
        "sheet_count": len(sheets),
        "sheets": sheets,
    }


def print_report(data: dict) -> None:
    """把结构化结果格式化为可读报告。"""
    print("=" * 72)
    print(f"模板结构报告  type={data.get('type')}  path={data.get('path')}")
    print("=" * 72)

    if data["type"] == "docx":
        print(f"\n标题数: {data['heading_count']}  表格数: {data['table_count']}")
        print("\n【标题大纲】")
        if not data["headings"]:
            print("  (未检测到 Heading/标题 样式；请结合段落样例判断结构)")
        for h in data["headings"]:
            indent = "  " * max(h["level"] - 1, 0)
            print(f"  {indent}[H{h['level']}|{h['style']}] {h['text']}")

        print("\n【表格】")
        if not data["tables"]:
            print("  (无表格)")
        for t in data["tables"]:
            print(f"  表{t['index']}: {t['rows']}行 x {t['cols']}列")
            print(f"    列头: {t['headers']}")
            if t["sample_row"]:
                print(f"    样例: {t['sample_row']}")

        print("\n【段落样例】（最多 40 条非空）")
        for p in data["paragraphs_sample"][:20]:
            print(f"  [{p['style']}] {p['text']}")

    elif data["type"] == "xlsx":
        print(f"\n工作表数: {data['sheet_count']}")
        for s in data["sheets"]:
            print(f"\n【Sheet】{s['name']}")
            print(f"  列头: {s['headers']}")
            for i, row in enumerate(s["sample_rows"], 1):
                print(f"  行{i}: {row}")

    # 机器可读尾部，便于后续脚本消费
    print("\n--- JSON ---")
    print(json.dumps(data, ensure_ascii=False, indent=2))


def main(argv: list[str]) -> int:
    _ensure_utf8_stdout()
    if len(argv) < 2:
        print(__doc__)
        print("用法: python inspect_docx.py <file.docx|file.xlsx>")
        return 1

    path = Path(argv[1]).expanduser().resolve()
    if not path.exists():
        print(f"错误: 文件不存在: {path}")
        return 1

    suffix = path.suffix.lower()
    if suffix == ".docx":
        data = inspect_docx(path)
    elif suffix in (".xlsx", ".xlsm"):
        data = inspect_xlsx(path)
    else:
        print(f"错误: 不支持的扩展名 {suffix}，仅支持 .docx / .xlsx")
        return 1

    print_report(data)
    return 0


if __name__ == "__main__":
    raise SystemExit(main(sys.argv))
